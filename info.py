from docx import Document

doc = Document("/Users/deyan/Downloads/KS3029 Keyestudio 4WD Mecanum Robot for Pico/1. Product Introduction/Product Introduction.docx")
parts = 3  # number of parts you want
total_paragraphs = len(doc.paragraphs)
chunk_size = total_paragraphs // parts

for i in range(parts):
    new_doc = Document()
    start = i * chunk_size
    end = (i + 1) * chunk_size if i < parts - 1 else total_paragraphs
    for p in doc.paragraphs[start:end]:
        new_doc.add_paragraph(p.text)
    new_doc.save(f"part_{i+1}.docx")
